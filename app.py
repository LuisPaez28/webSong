import streamlit as st
import librosa
import numpy as np
import whisper
import os
import subprocess
import imageio_ffmpeg
from docx import Document
from io import BytesIO

# Configuración de página
st.set_page_config(page_title="Music Recovery Web", layout="wide")

CHORD_NAMES = ['C', 'C#', 'D', 'D#', 'E', 'F', 'F#', 'G', 'G#', 'A', 'A#', 'B']

def get_chord(chroma):
    idx = np.argmax(chroma)
    suffix = "" if chroma[(idx + 4) % 12] > chroma[(idx + 3) % 12] else "m"
    return f"{CHORD_NAMES[idx]}{suffix}"

st.title("🎸 Music Recovery Studio (Web Local)")
st.markdown("Sube tu video o audio para recuperar la letra y los acordes de forma privada.")

# --- BARRA LATERAL ---
st.sidebar.header("Configuración")
modelo = st.sidebar.selectbox("Modelo de Whisper", ["base", "tiny", "small"], index=0)

# --- CARGA DE ARCHIVOS ---
uploaded_file = st.file_uploader("Elige un archivo de video o audio", type=["mp4", "mov", "wav", "mp3"])

if uploaded_file is not None:
    # Guardar archivo temporalmente
    temp_input = "temp_input" + os.path.splitext(uploaded_file.name)[1]
    with open(temp_input, "wb") as f:
        f.write(uploaded_file.getbuffer())

    if st.button("Comenzar Análisis"):
        with st.spinner("Procesando... esto puede tardar dependiendo de la duración."):
            
            # 1. Convertir a WAV si es video
            audio_path = "temp_audio.wav"
            ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
            subprocess.run([ffmpeg_exe, "-y", "-i", temp_input, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", audio_path], 
                           capture_output=True)

            # 2. IA Whisper
            st.info("Escuchando letra...")
            model = whisper.load_model(modelo, device="cpu")
            result = model.transcribe(audio_path, fp16=False)

            # 3. Análisis de Acordes
            st.info("Analizando armonía...")
            y, sr = librosa.load(audio_path, sr=None)
            chroma = librosa.feature.chroma_cqt(y=y, sr=sr, hop_length=1024)
            times = librosa.times_like(chroma, sr=sr, hop_length=1024)

            # 4. Mostrar Resultados y Generar Word
            doc = Document()
            doc.add_heading('Cancionero Recuperado', 0)
            
            st.subheader("📝 Resultados")
            col1, col2 = st.columns([1, 2])
            
            resultados_txt = ""
            for seg in result['segments']:
                idx_start = np.searchsorted(times, seg['start'])
                idx_end = np.searchsorted(times, seg['end'])
                
                # Lógica de homologación (máximo 4 acordes)
                paso = max(1, (idx_end - idx_start) // 4)
                acordes_list = []
                ultimo = ""
                for i in range(idx_start, idx_end, paso):
                    act = get_chord(chroma[:, i])
                    if act != ultimo:
                        acordes_list.append(act)
                        ultimo = act
                
                acordes_str = "   ".join(acordes_list)
                tiempo = f"[{int(seg['start']//60):02d}:{int(seg['start']%60):02d}]"
                texto = seg['text'].strip()

                # Mostrar en la Web
                with st.container():
                    st.markdown(f"**{tiempo} {acordes_str}**")
                    st.write(texto)
                    st.divider()

                # Agregar al documento Word
                p = doc.add_paragraph()
                run = p.add_run(f"{tiempo}  {acordes_str}")
                run.bold = True
                doc.add_paragraph(texto)

            # 5. Botón de Descarga Word
            buffer = BytesIO()
            doc.save(buffer)
            st.download_button(
                label="📄 Descargar cancionero en Word",
                data=buffer.getvalue(),
                file_name="cancionero.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Limpieza
    if os.path.exists(temp_input): os.remove(temp_input)
    if os.path.exists("temp_audio.wav"): os.remove("temp_audio.wav")