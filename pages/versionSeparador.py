import streamlit as st
import librosa
import numpy as np
import whisper
import os
import subprocess
import imageio_ffmpeg
from docx import Document
from io import BytesIO
from spleeter.separator import Separator # <--- LA NUEVA MAGIA

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Music Recovery Studio Web", layout="wide", page_icon="🎸")

# --- CACHÉ DE MODELOS ---
@st.cache_resource
def load_whisper_model(model_size):
    try:
        return whisper.load_model(model_size, device="cpu")
    except:
        return whisper.load_model(model_size, device="cpu", download_root="./models")

@st.cache_resource
def load_spleeter():
    """Carga el modelo de separación de Spleeter en memoria"""
    # 'spleeter:2stems' significa separar en 2 pistas: Voz + Acompañamiento
    return Separator('spleeter:2stems')

CHORD_NAMES = ['C', 'C#', 'D', 'D#', 'E', 'F', 'F#', 'G', 'G#', 'A', 'A#', 'B']

def get_chord(chroma):
    idx = np.argmax(chroma)
    suffix = "" if chroma[(idx + 4) % 12] > chroma[(idx + 3) % 12] else "m"
    return f"{CHORD_NAMES[idx]}{suffix}"

# --- INTERFAZ ---
st.title("🎸 Music Recovery Studio v2.0")

# TRES PESTAÑAS AHORA
tab_conv, tab_ana, tab_split = st.tabs([
    "1. Convertidor Video a Audio", 
    "2. Analizador de Canciones", 
    "3. Separador (Karaoke)"
])

# --- PESTAÑA 1: CONVERTIDOR (Igual que antes) ---
with tab_conv:
    st.header("Extraer Audio de Video")
    video_file = st.file_uploader("Sube tu video", type=["mp4", "mov", "avi", "mkv"])
    if video_file and st.button("Extraer Audio"):
        with st.spinner("Procesando con FFmpeg..."):
            temp_vid = "temp_video" + os.path.splitext(video_file.name)[1]
            with open(temp_vid, "wb") as f: f.write(video_file.getbuffer())
            output_audio = "audio_extraido.wav"
            subprocess.run([imageio_ffmpeg.get_ffmpeg_exe(), "-y", "-i", temp_vid, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", output_audio], capture_output=True)
            with open(output_audio, "rb") as file:
                st.download_button("🎵 Descargar .WAV", file, file_name="audio.wav", mime="audio/wav")
            os.remove(temp_vid); os.remove(output_audio)

# --- PESTAÑA 2: ANALIZADOR (Igual que antes) ---
with tab_ana:
    st.header("Recuperar Letra y Acordes")
    audio_ana = st.file_uploader("Sube audio para acordes", type=["wav", "mp3"], key="ana")
    mod_size = st.selectbox("Modelo IA", ["base", "tiny"], key="whisper_model")
    
    if audio_ana and st.button("Analizar Canción"):
        with st.spinner("Transcribiendo y analizando armonía..."):
            model = load_whisper_model(mod_size)
            temp_ana = "temp_ana.wav"
            with open(temp_ana, "wb") as f: f.write(audio_ana.getbuffer())
            
            # Whisper + Librosa
            result = model.transcribe(temp_ana, fp16=False)
            y, sr = librosa.load(temp_ana, sr=None)
            chroma = librosa.feature.chroma_cqt(y=y, sr=sr, hop_length=1024)
            times = librosa.times_like(chroma, sr=sr, hop_length=1024)
            
            doc = Document(); doc.add_heading('Cancionero', 0)
            st.markdown("---")
            for seg in result['segments']:
                idx_s = np.searchsorted(times, seg['start'])
                idx_e = np.searchsorted(times, seg['end'])
                paso = max(1, (idx_e - idx_s) // 4)
                acordes = []
                last = ""
                for i in range(idx_s, idx_e, paso):
                    c = get_chord(chroma[:, i])
                    if c != last: acordes.append(c); last = c
                
                linea = "   ".join(acordes)
                t = f"[{int(seg['start']//60):02d}:{int(seg['start']%60):02d}]"
                st.markdown(f"**{t} {linea}**"); st.write(seg['text'])
                p = doc.add_paragraph(); r = p.add_run(f"{t}  {linea}"); r.bold = True
                doc.add_paragraph(seg['text'])
            
            buf = BytesIO(); doc.save(buf)
            st.download_button("📄 Descargar Word", buf.getvalue(), "cancionero.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            os.remove(temp_ana)

# --- PESTAÑA 3: SEPARADOR (NUEVA) ---
with tab_split:
    st.header("Separador de Pistas (IA)")
    st.markdown("Separa la voz de la música para crear pistas de karaoke o practicar.")
    
    audio_split = st.file_uploader("Sube una canción", type=["wav", "mp3"], key="split")
    
    if audio_split and st.button("Separar Voz y Música"):
        with st.spinner("La IA está separando las pistas (esto usa mucha CPU)..."):
            # 1. Guardar archivo original
            temp_split = "temp_split_source" + os.path.splitext(audio_split.name)[1]
            with open(temp_split, "wb") as f: f.write(audio_split.getbuffer())
            
            # 2. Cargar Spleeter y Procesar
            separator = load_spleeter()
            output_dir = "output_split"
            
            # Esta línea hace toda la magia
            separator.separate_to_file(temp_split, output_dir)
            
            # 3. Localizar los archivos generados
            # Spleeter crea una carpeta con el nombre del archivo original (sin extensión)
            base_name = os.path.splitext(temp_split)[0]
            # Nota: Spleeter a veces normaliza nombres, así que buscamos en la carpeta
            generated_folder = os.path.join(output_dir, base_name)
            
            vocals_path = os.path.join(generated_folder, "vocals.wav")
            music_path = os.path.join(generated_folder, "accompaniment.wav")
            
            col1, col2 = st.columns(2)
            
            # 4. Mostrar Resultados
            with col1:
                st.success("🎤 Voz Aislada")
                if os.path.exists(vocals_path):
                    st.audio(vocals_path)
                    with open(vocals_path, "rb") as f:
                        st.download_button("Descargar Voz", f, "voz_aislada.wav", "audio/wav")
            
            with col2:
                st.success("🎹 Música (Instrumental)")
                if os.path.exists(music_path):
                    st.audio(music_path)
                    with open(music_path, "rb") as f:
                        st.download_button("Descargar Pista", f, "pista_instrumental.wav", "audio/wav")

            # Limpieza (Opcional, cuidado de no borrar antes de descargar)
            # os.remove(temp_split)