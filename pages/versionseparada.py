import streamlit as st
import librosa
import numpy as np
import whisper
import os
import subprocess
import imageio_ffmpeg
from docx import Document
from io import BytesIO

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Traductor de Canciones con convertidor de video a audio", layout="wide", page_icon="🎸")

# --- CACHÉ DE MODELO ---
@st.cache_resource
def load_whisper_model(model_size):
    return whisper.load_model(model_size, device="cpu")

CHORD_NAMES = ['C', 'C#', 'D', 'D#', 'E', 'F', 'F#', 'G', 'G#', 'A', 'A#', 'B']

def get_chord(chroma):
    idx = np.argmax(chroma)
    suffix = "" if chroma[(idx + 4) % 12] > chroma[(idx + 3) % 12] else "m"
    return f"{CHORD_NAMES[idx]}{suffix}"

# --- INTERFAZ DE USUARIO ---
st.title("Recuperador de acordes y letras de canciones con Convertidor de Video a Audio")

# Creamos las pestañas igual que en la versión nativa
tab_conv, tab_ana = st.tabs(["1. Convertidor Video a Audio", "2. Analizador de Canciones"])

# --- PESTAÑA 1: CONVERTIDOR ---
with tab_conv:
    st.header("Extraer Audio de Video")
    video_file = st.file_uploader("Sube tu video para extraer el sonido", type=["mp4", "mov", "avi", "mkv"])
    
    if video_file:
        if st.button("Procesar y Extraer Audio"):
            with st.spinner("Extrayendo audio con FFmpeg..."):
                temp_vid = "temp_video_conv" + os.path.splitext(video_file.name)[1]
                with open(temp_vid, "wb") as f:
                    f.write(video_file.getbuffer())
                
                output_audio = "audio_extraido.wav"
                ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
                
                # Comando FFmpeg optimizado para música (Mono, 16kHz)
                subprocess.run([ffmpeg_exe, "-y", "-i", temp_vid, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", output_audio], capture_output=True)
                
                # Botón de descarga para el audio generado
                with open(output_audio, "rb") as file:
                    st.success("¡Audio extraído con éxito!")
                    st.download_button(
                        label="🎵 Descargar archivo .WAV",
                        data=file,
                        file_name=os.path.splitext(video_file.name)[0] + ".wav",
                        mime="audio/wav"
                    )
                
                # Limpieza
                os.remove(temp_vid)
                os.remove(output_audio)

# --- PESTAÑA 2: ANALIZADOR ---
with tab_ana:
    st.header("Recuperar Letra y Acordes")
    audio_file = st.file_uploader("Sube el audio (.wav o .mp3) para analizar", type=["wav", "mp3"])
    
    modelo_size = st.selectbox("Calidad de la IA (Whisper)", ["base", "tiny", "small"])

    if audio_file:
        if st.button("Iniciar Análisis Completo"):
            with st.spinner("Analizando... esto tomará un momento."):
                # Cargar modelo
                model = load_whisper_model(modelo_size)
                
                # Guardar temporal
                temp_audio = "temp_ana.wav"
                with open(temp_audio, "wb") as f:
                    f.write(audio_file.getbuffer())

                # 1. Transcribir letra
                st.info("Transcribiendo letra...")
                result = model.transcribe(temp_audio, fp16=False)

                # 2. Analizar acordes
                st.info("Analizando acordes homologados...")
                y, sr = librosa.load(temp_audio, sr=None)
                chroma = librosa.feature.chroma_cqt(y=y, sr=sr, hop_length=1024)
                times = librosa.times_like(chroma, sr=sr, hop_length=1024)

                # 3. Mostrar y preparar Word
                doc = Document()
                doc.add_heading('Cancionero Recuperado', 0)
                
                st.markdown("---")
                for seg in result['segments']:
                    # Lógica de búsqueda de acordes por verso
                    idx_start = np.searchsorted(times, seg['start'])
                    idx_end = np.searchsorted(times, seg['end'])
                    
                    paso = max(1, (idx_end - idx_start) // 4)
                    acordes_v = []
                    last = ""
                    for i in range(idx_start, idx_end, paso):
                        c = get_chord(chroma[:, i])
                        if c != last:
                            acordes_v.append(c)
                            last = c
                    
                    acordes_linea = "   ".join(acordes_v)
                    tiempo = f"[{int(seg['start']//60):02d}:{int(seg['start']%60):02d}]"
                    texto = seg['text'].strip()

                    # Mostrar en web
                    st.markdown(f"#### {tiempo} &nbsp;&nbsp; **{acordes_linea}**")
                    st.write(texto)
                    
                    # Añadir a Word
                    p = doc.add_paragraph()
                    run = p.add_run(f"{tiempo}  {acordes_linea}")
                    run.bold = True
                    doc.add_paragraph(texto)

                # Botón Word
                buffer = BytesIO()
                doc.save(buffer)
                st.download_button(
                    label="📄 Descargar (.docx)",
                    data=buffer.getvalue(),
                    file_name="mi_cancionero.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                os.remove(temp_audio)